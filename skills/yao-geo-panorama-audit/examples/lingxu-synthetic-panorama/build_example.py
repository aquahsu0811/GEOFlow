# Copyright © 2026 姚金刚. All rights reserved.
# Project: yao-geo-panorama-audit
# Created by: 姚金刚
# Date: 2026-05-16
# X: https://x.com/yaojingang

from __future__ import annotations

import html
import shutil
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


BASE = Path(__file__).resolve().parent
SLUG = "lingxu-panorama-audit"
MD_PATH = BASE / f"{SLUG}.md"
HTML_PATH = BASE / f"{SLUG}.html"
DOCX_PATH = BASE / f"{SLUG}.docx"
PDF_PATH = BASE / f"{SLUG}.pdf"
REVIEW_PATH = BASE / "review-notes.md"
DOCX_SAFE_TABLE_WIDTH_CM = 15.8
COPYRIGHT_COMMENT = """<!--
Copyright © 2026 姚金刚. All rights reserved.
Project: yao-geo-panorama-audit
Created by: 姚金刚
Date: 2026-05-16
X: https://x.com/yaojingang
-->

"""


REPORT = {
    "title": "GEO 全景诊断与机会地图示例报告",
    "brand": "岭序商机云",
    "subtitle": "官网抓取与公开事实交叉验证合成样例",
    "example_nature": "官网抓取与公开事实交叉验证合成样例",
    "notice": "本报告为 yao-geo-panorama-audit 的合成示例，用于展示官网抓取、断言级交叉验证、站内方案和站外方案的报告结构与四格式版式，不代表对任何真实品牌的商业结论。",
    "meta": [
        ("品牌", "岭序商机云（合成 B2B 线索管理软件）"),
        ("官网", "https://example.com/lingxu"),
        ("行业与区域", "企业服务 / 中国大陆"),
        ("目标用户", "中小企业销售负责人、市场负责人、增长负责人"),
        ("官网抓取范围", "首页、产品页、价格页、帮助中心、案例页、博客、关于页、sitemap"),
        ("交叉验证范围", "官方帮助文档、客户授权材料、公众号、媒体稿、行业报告、竞品公开页"),
        ("诊断方法", "官网优先；断言级来源台账；站内/站外方案拆分"),
        ("版本日期", "2026-05-21"),
    ],
    "summary": [
        ("官网事实不够原子化", "首页能表达线索管理定位，但缺少可直接抽取的品牌定义、产品边界、适用对象和不适用场景。"),
        ("价格与案例证据不足", "价格页没有清晰套餐边界，案例页缺少授权状态、行业、问题、方案和结果字段。"),
        ("九类用户问题覆盖不均", "教程、价格、风险和真实性问题缺少站内答案，容易让外部解释替代官方口径。"),
        ("外部信源薄弱", "公众号、媒体、百科、社区和行业报告缺少可交叉验证的公开材料，权威信号不足。"),
        ("P0 分成站内和站外", "站内先做实体事实卡、产品页和价格页重构；站外先补客户案例、媒体稿和公众号结构化证据。"),
    ],
    "method_basis": [
        ["官网抓取", "从导航、页脚、sitemap 和核心页面建立页面清单", "URL、页面类型、标题、可抓取性、更新时间", "避免只凭主观印象诊断"],
        ["断言抽取", "把页面内容拆成品牌、产品、价格、案例、资质、限制等事实", "断言、来源、核验状态、用途", "避免大段营销话术无法引用"],
        ["交叉验证", "用官方二级来源和公开第三方来源核验关键事实", "交叉验证来源、冲突、待确认项", "避免把未授权材料写成公开事实"],
        ["站内/站外拆分", "把页面结构修复和外部证据建设分开排优先级", "负责人、工作量、验收标准", "避免所有建议混成泛内容建设"],
    ],
    "authority_map": [
        ["GEO 研究", "内容需要面向生成式答案中的可见性和引用做优化", "用于设定 GEO 特征评分和机会地图", "不直接承诺平台推荐概率"],
        ["生成式搜索可验证性研究", "流畅答案可能存在未支持断言和错配引用", "用于要求断言级来源台账和交叉验证", "不把单个来源数量等同可信度"],
        ["Google 实用可靠内容", "内容应完整、原创、可信、服务真实用户目标", "用于补充完整性、专业性和用户问题覆盖", "不为搜索引擎堆砌内容"],
        ["Google 结构化数据规范", "标记内容必须与页面可见内容一致且保持最新", "用于约束 Schema 建议必须绑定页面事实", "不对不可见内容加结构化标记"],
        ["Schema.org", "Organization、Product/Service、FAQ、Breadcrumb 等可表达实体和关系", "用于站内实体和页面结构设计", "不把 Schema 当作排名保证"],
        ["W3C/MDN 导航规范", "长页面需要一致导航，sticky 菜单可在滚动中保持可见", "用于 HTML 报告固定菜单和锚点设计", "不遮挡正文和表格"],
    ],
    "question_coverage": [
        ["推荐", "适合中小企业的线索管理软件有哪些？", "部分覆盖", "缺少品牌适用场景短答案", "新增选型入口页"],
        ["比较", "岭序商机云和脉冲 CRM 有什么区别？", "弱覆盖", "无竞品对比和边界说明", "新增对比页"],
        ["替代", "如果不用岭序商机云，有哪些替代方案？", "未覆盖", "无法解释替代场景", "新增替代方案 FAQ"],
        ["教程", "销售团队如何从表格迁移到线索管理工具？", "部分覆盖", "教程分散且缺少步骤索引", "重构帮助中心入口"],
        ["价格", "线索管理系统一般怎么收费？岭序商机云价格是否公开？", "弱覆盖", "套餐、限制和咨询边界不清", "重构价格页"],
        ["风险", "中小企业上线线索管理系统有哪些风险？", "未覆盖", "缺少上线条件、数据迁移和权限风险", "新增风险说明页"],
        ["真实性", "岭序商机云是真的吗？有哪些公开案例？", "弱覆盖", "公司主体和案例授权不足", "完善关于页和案例页"],
        ["购买决策", "什么情况下应该选择岭序商机云？", "部分覆盖", "适用/不适用场景不完整", "新增采购决策 FAQ"],
        ["场景解决", "教育培训机构如何管理销售线索并提升跟进效率？", "未覆盖", "行业页缺失", "新增教育培训行业页"],
    ],
    "baseline": [
        ["品牌定义", "首页 / 关于页", "面向中小企业的线索管理工具", "产品手册 / 公司主体资料", "部分核验", "改成 80 字实体事实卡并标注更新时间"],
        ["产品边界", "产品页 / 帮助中心", "线索录入、分配、跟进、转化分析", "帮助文档 / 销售资料", "已核验", "增加功能边界表和不适用场景"],
        ["价格套餐", "价格页", "提供咨询制报价", "销售资料", "待确认", "补套餐、计费口径、限制和更新时间"],
        ["客户案例", "案例页", "覆盖教育培训和企业服务客户", "客户授权材料", "待确认", "补行业、规模、问题、方案、结果和授权状态"],
        ["资质与主体", "关于页 / 页脚", "公司主体、联系方式、服务地域", "工商信息 / 备案信息", "部分核验", "补主体信息和服务边界"],
        ["内容资产", "博客 / 帮助中心", "有教程文章但缺少问题入口", "站内搜索 / 客服问答", "部分核验", "按九类问题重组专题"],
    ],
    "crawl_inventory": [
        ["首页", "https://example.com/lingxu", "可访问", "品牌定位、产品入口", "缺少更新时间和来源区", "P0 补实体事实卡"],
        ["产品页", "/product", "可访问", "功能模块、流程说明", "缺少功能边界表", "P0 重构产品结构"],
        ["价格页", "/pricing", "内容不足", "咨询制报价", "套餐、限制、更新时间缺失", "P0 重构商业信息"],
        ["案例页", "/customers", "内容不足", "客户故事摘要", "授权状态和结果字段缺失", "P1 建案例模板"],
        ["帮助中心", "/help", "可访问", "教程文章", "缺少问题索引和内链", "P1 建教程中心"],
        ["sitemap", "/sitemap.xml", "待确认", "页面发现", "未确认是否覆盖帮助中心", "P1 技术检查"],
    ],
    "entity_profile": [
        ["品牌名", "岭序商机云", "官网首页", "已确认", "需要固定中文定义"],
        ["英文名/别名", "Lingxu Lead Cloud（示例）", "品牌资料", "待确认", "避免同名或机器翻译混用"],
        ["公司主体", "岭序科技（示例）", "关于页 / 备案", "待确认", "补 legalName、备案和联系方式"],
        ["产品关系", "线索管理、跟进分配、转化分析", "产品页 / 帮助中心", "部分确认", "补产品服务地图"],
        ["服务地域", "中国大陆中小企业", "销售资料", "待确认", "补服务边界和不适用场景"],
    ],
    "product_map": [
        ["线索管理", "线索录入、分配、去重、标签", "市场与销售负责人", "咨询制", "产品页 / 帮助中心", "补功能边界"],
        ["销售跟进", "任务提醒、跟进记录、阶段推进", "销售团队主管", "咨询制", "帮助中心", "补教程入口"],
        ["转化分析", "渠道、转化率、跟进效率", "增长负责人", "咨询制", "产品页 / 报表页", "补指标定义"],
        ["行业方案", "教育培训、企业服务", "行业负责人", "未公开", "案例页", "补行业页和案例证据"],
    ],
    "technical_audit": [
        ["sitemap", "待确认", "未确认是否包含帮助中心和案例页", "页面发现不完整", "生成并提交完整 sitemap"],
        ["robots", "待确认", "未检查禁止抓取目录", "核心页面可能不可访问", "检查 robots 和重要目录"],
        ["canonical", "弱", "价格、案例可能存在重复入口", "重复页面稀释实体信号", "补 canonical 和规范 URL"],
        ["结构化数据", "缺失", "未见 Organization、Product/Service、FAQ", "实体和问答难以机器读取", "按可见内容补 JSON-LD"],
        ["PDF/图片文本", "弱", "部分销售材料可能为图片或 PDF", "关键事实不可抽取", "改成 HTML 文本或提供可复制摘要"],
    ],
    "competitors": [
        ["脉冲 CRM", "案例页字段完整，媒体稿较多", "岭序案例授权和结果字段不足", "客户故事模板、行业页结构", "不能直接复制其大客户口径"],
        ["销售易", "行业方案和渠道材料丰富", "岭序行业页缺失", "行业问题入口和解决路径", "需保持中小企业轻量定位"],
        ["简道云", "教程、社区和低代码场景内容多", "岭序帮助中心缺少索引", "教程入口、步骤化 FAQ", "避免被误归类为低代码平台"],
        ["岭序商机云", "产品说明较清楚，帮助文档可扩展", "外部信源和实体页薄弱", "先补官网事实，再补外部证据", "保持合成样例口径"],
    ],
    "scores": [
        ["语义密度", "2.5/5", "官网反复出现线索、客户、销售，但缺少场景化同义表达", "补行业场景页和问题页"],
        ["结构规范性", "2.0/5", "首页和产品页有段落，但缺少摘要、表格、常见问题和定义区", "重构核心页面结构"],
        ["可引用性", "1.8/5", "公开来源少，页面缺少可直接引用的短句与数据", "新增事实卡片和来源台账"],
        ["权威信号", "1.6/5", "客户案例、资质、媒体报道和第三方评价不足", "补强外部证据"],
        ["可读性", "3.2/5", "基础表达清楚，但页面信息层级不够稳定", "压缩长段落并增加表格"],
        ["鲁棒性", "2.1/5", "同义问法下官网答案入口不稳定", "建立品牌实体档案"],
        ["新颖性", "2.8/5", "有中小企业场景，但差异化表达未形成术语", "定义核心概念"],
        ["跨域贡献", "1.5/5", "公众号、视频号、社区、行业报告覆盖不足", "建设外部信源矩阵"],
    ],
    "opportunities": [
        ["站内", "品牌实体事实卡", "官网定义不可直接引用", "高", "低", "P0", "内容负责人", "首页和关于页均出现统一定义、别名和适用边界"],
        ["站内", "价格页重构", "套餐和限制不清", "高", "中", "P0", "产品市场", "价格页包含计费口径、咨询边界和更新时间"],
        ["站内", "九类问题 FAQ", "高频问题缺少站内答案", "高", "中", "P0", "GEO 顾问", "覆盖 30 个问题并接入内链"],
        ["站外", "客户案例公开证据", "行业案例不可核验", "中", "中", "P1", "销售运营", "至少 6 个案例有授权状态和结果字段"],
        ["站外", "公众号与媒体稿", "中文外部信源不足", "中", "中", "P1", "品牌负责人", "发布 4 篇案例文章和 2 篇媒体稿"],
        ["站外", "行业报告或社区问答", "场景问题缺少第三方解释", "中", "中", "P2", "增长负责人", "形成 3 个可公开引用的行业观点来源"],
    ],
    "repairs": [
        ["官网首页", "品牌定义没有固定短句", "新增 80 字品牌定义、别名、适用/不适用场景", "实体事实卡、来源区", "品牌资料、产品手册", "官网能复制出统一品牌定义"],
        ["产品页", "功能描述偏营销，缺少结构化事实", "增加功能表、流程图、限制条件和更新时间", "功能边界表", "产品文档、帮助中心", "产品边界可被表格化引用"],
        ["价格页", "价格与套餐边界不清", "增加套餐对比表、适用对象和咨询说明", "价格 FAQ、更新时间", "销售资料、官网价格页", "价格问题不再需要外部猜测"],
        ["帮助中心", "教程分散，缺少入口页", "建立“从线索到成交”教程索引", "教程中心页", "帮助文档、客服问答", "教程类问题能定位到具体步骤"],
        ["案例页", "案例内容少且不可结构化引用", "发布案例模板：背景、问题、方案、结果、限制", "案例事实卡", "客户授权、销售记录", "公开案例字段完整且授权清楚"],
    ],
    "offsite_plan": [
        ["公众号", "发布客户案例和选型 FAQ", "品牌真实性、场景解决、购买决策", "P1", "品牌负责人", "4 篇文章均包含来源、客户授权状态和官网链接"],
        ["媒体稿", "发布产品边界和行业场景稿", "权威信号、产品边界", "P1", "公关负责人", "至少 2 篇媒体稿可支撑品牌定义和行业定位"],
        ["百科/知识库", "补全品牌主体、产品和官网入口", "实体消歧、真实性", "P2", "运营负责人", "百科条目与官网事实一致且无夸大表述"],
        ["合作伙伴页", "补合作伙伴解决方案和集成说明", "渠道、服务地域、生态", "P2", "渠道负责人", "合作伙伴页能回链到官网对应方案"],
        ["社区问答", "沉淀上线风险、数据迁移、权限管理问答", "风险、教程、场景解决", "P2", "客户成功", "公开问答不泄露客户隐私且可链接官网教程"],
    ],
    "sources": [
        ["品牌定义", "官网与产品手册", "部分核验", "建立实体档案", "合成样例中仅展示字段结构"],
        ["价格与套餐", "销售资料与价格页", "待确认", "修复价格类问题", "正式项目需标注更新时间"],
        ["客户案例", "客户授权与销售记录", "待确认", "补强权威信号", "未授权案例不得写成公开事实"],
        ["外部信源", "公众号、媒体、百科、社区、视频号", "缺口明显", "判断跨域贡献", "优先补可公开引用的中文来源"],
    ],
    "risk_register": [
        ["客户案例未授权", "中", "P0", "销售运营", "未授权客户不得进入公开页面，先做匿名结构化案例"],
        ["价格口径过期", "高", "P0", "产品市场", "价格页必须标注更新时间和咨询边界"],
        ["品牌主体不清", "高", "P0", "品牌负责人", "补公司主体、官网、备案和联系方式"],
        ["外部信源不足", "中", "P1", "公关负责人", "优先补媒体稿和公众号案例"],
        ["技术抓取不完整", "中", "P1", "技术负责人", "检查 sitemap、robots、canonical 和 JS 渲染"],
    ],
    "diagnostic_questions": [
        "适合中小企业的线索管理软件有哪些？",
        "岭序商机云和脉冲 CRM 有什么区别？",
        "线索管理系统一般怎么收费？",
        "销售团队如何从表格迁移到线索管理工具？",
        "岭序商机云是真的吗？有哪些公开案例？",
        "如果不用岭序商机云，有哪些替代方案？",
        "中小企业上线线索管理系统有哪些风险？",
        "什么情况下应该选择岭序商机云？",
        "教育培训机构如何管理销售线索并提升跟进效率？",
    ],
    "section_meanings": {
        "method_basis": "本报告把 GEO 诊断视为官网事实源整理、断言级核验、公开来源交叉验证和站内/站外证据建设。",
        "authority_map": "权威参考用于约束诊断口径：既吸收 GEO 和可验证性研究，也借鉴内容质量、结构化数据和可访问导航规范。",
        "question_coverage": "用户问题覆盖矩阵用于检查官网是否能回答九类高频问题，缺口会直接转化为站内页面或站外证据任务。",
        "baseline": "岭序商机云的核心问题是关键断言不够可抽取，价格、案例和主体信息仍需交叉验证。",
        "crawl_inventory": "抓取清单把“有没有页面”进一步拆成可访问、可抽取、缺失字段和下一步动作。",
        "entity_profile": "实体档案用于解决品牌、公司主体、产品关系和服务地域的一致性问题。",
        "product_map": "产品服务地图把功能、对象、商业口径和来源绑定，避免产品边界被泛化。",
        "technical_audit": "技术诊断保证内容不仅写出来，还能被稳定发现、抓取、解析和结构化表达。",
        "competitors": "竞品优势不只来自官网页面数量，也来自媒体、百科、案例和社区讨论等外部信源的共同作用。",
        "scores": "当前最短路径不是扩大泛内容产量，而是先修复可引用事实、品牌实体和核心页面结构。",
        "opportunities": "P0 机会集中在站内实体事实、价格页和九类问题 FAQ；站外证据按客户授权和传播资源逐步推进。",
        "risk_register": "风险台账把无法确认、过期、授权、技术和资源问题显式化，避免报告把假设写成事实。",
    },
    "conclusion": "岭序商机云的 GEO 基线问题主要不是“品牌完全没有内容”，而是官网事实不够原子化、外部信源不够稳定、核心页面缺少可引用结构。建议先完成站内 P0 修复，再补客户案例、媒体稿和公众号证据，并按季度复查官网来源和外部证据变化。",
}


HEADERS = {
    "method_basis": ["方法环节", "执行方式", "关键字段", "控制风险"],
    "authority_map": ["参考来源", "采用原则", "报告落点", "边界"],
    "question_coverage": ["意图类型", "用户问题", "官网覆盖状态", "缺口判断", "站内用途"],
    "baseline": ["事实域", "官网页面/入口", "官网可抽取事实", "交叉验证来源", "核验状态", "优化方向"],
    "crawl_inventory": ["页面类型", "URL/入口", "状态", "可抽取事实", "缺失字段", "下一步"],
    "entity_profile": ["实体项", "当前口径", "来源", "核验状态", "处理建议"],
    "product_map": ["产品/模块", "核心能力", "适用对象", "商业口径", "事实来源", "缺口"],
    "technical_audit": ["技术项", "状态", "证据", "风险", "修复动作"],
    "competitors": ["对标对象", "公开可见优势", "目标品牌差距", "可借鉴证据", "注意边界"],
    "scores": ["特征", "评分", "证据判断", "优先动作"],
    "opportunities": ["方案类型", "机会", "对应问题", "价值", "工作量", "优先级", "负责人建议", "验收标准"],
    "repairs": ["页面或模块", "主要问题", "修复动作", "输出组件", "依赖来源", "验收口径"],
    "offsite_plan": ["渠道/信源", "建设动作", "支撑断言", "优先级", "负责人建议", "验收口径"],
    "sources": ["事实项", "来源类型", "核验状态", "用途", "备注"],
    "risk_register": ["风险", "影响", "优先级", "负责人建议", "处理方式"],
}


def markdown_table(headers: list[str], rows: list[list[str]]) -> str:
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join(["---"] * len(headers)) + " |",
    ]
    for row in rows:
        safe_row = [str(cell).replace("\n", "<br>") for cell in row]
        lines.append("| " + " | ".join(safe_row) + " |")
    return "\n".join(lines)


def render_markdown() -> str:
    meta = "\n".join([f"- {key}：{value}" for key, value in REPORT["meta"]])
    summary = "\n".join([f"- **{title}**：{detail}" for title, detail in REPORT["summary"]])
    questions = "\n".join([f"{idx}. {question}" for idx, question in enumerate(REPORT["diagnostic_questions"], 1)])
    meanings = REPORT["section_meanings"]
    return COPYRIGHT_COMMENT + f"""# {REPORT["title"]}

**{REPORT["brand"]}｜{REPORT["subtitle"]}**

> {REPORT["notice"]}

## 诊断元信息

{meta}

## 执行摘要

{summary}

## 方法依据与评分口径

{markdown_table(HEADERS["method_basis"], REPORT["method_basis"])}

判断含义：{meanings["method_basis"]}

## 权威参考映射

{markdown_table(HEADERS["authority_map"], REPORT["authority_map"])}

判断含义：{meanings["authority_map"]}

## 用户问题覆盖矩阵

{markdown_table(HEADERS["question_coverage"], REPORT["question_coverage"])}

判断含义：{meanings["question_coverage"]}

## 官网抓取与事实交叉验证

{markdown_table(HEADERS["baseline"], REPORT["baseline"])}

判断含义：{meanings["baseline"]}

## 官网抓取覆盖清单

{markdown_table(HEADERS["crawl_inventory"], REPORT["crawl_inventory"])}

判断含义：{meanings["crawl_inventory"]}

## 品牌实体与事实档案

{markdown_table(HEADERS["entity_profile"], REPORT["entity_profile"])}

判断含义：{meanings["entity_profile"]}

## 产品服务与商业信息地图

{markdown_table(HEADERS["product_map"], REPORT["product_map"])}

判断含义：{meanings["product_map"]}

## 技术可抓取与结构化数据诊断

{markdown_table(HEADERS["technical_audit"], REPORT["technical_audit"])}

判断含义：{meanings["technical_audit"]}

## 竞品与外部信源对标

{markdown_table(HEADERS["competitors"], REPORT["competitors"])}

判断含义：{meanings["competitors"]}

## GEO 特征评分

{markdown_table(HEADERS["scores"], REPORT["scores"])}

判断含义：{meanings["scores"]}

## 机会地图与优先级矩阵

{markdown_table(HEADERS["opportunities"], REPORT["opportunities"])}

判断含义：{meanings["opportunities"]}

## 站内系统方案

{markdown_table(HEADERS["repairs"], REPORT["repairs"])}

## 站外证据建设方案

{markdown_table(HEADERS["offsite_plan"], REPORT["offsite_plan"])}

## 诊断问题池

{questions}

## 来源台账与待确认项

{markdown_table(HEADERS["sources"], REPORT["sources"])}

## 风险、假设与待确认项

{markdown_table(HEADERS["risk_register"], REPORT["risk_register"])}

判断含义：{meanings["risk_register"]}

## 结论

{REPORT["conclusion"]}
"""


def render_html() -> str:
    css = """
:root {
  color-scheme: light;
  --ink: #18212f;
  --muted: #526070;
  --line: #d8e0e8;
  --soft: #f5f7fa;
  --accent: #1d5d8c;
}
* { box-sizing: border-box; }
html, body { margin: 0; padding: 0; background: #ffffff; color: var(--ink); }
html { scroll-behavior: smooth; }
body {
  font-family: -apple-system, BlinkMacSystemFont, "PingFang SC", "Hiragino Sans GB", "Microsoft YaHei", sans-serif;
  font-size: 15px;
  line-height: 1.66;
}
.page {
  max-width: 1060px;
  margin: 0 auto;
  padding: 46px 42px 64px;
  background: #ffffff;
}
.cover {
  border-bottom: 2px solid var(--ink);
  padding-bottom: 20px;
  margin-bottom: 26px;
}
h1, h2, h3 { margin: 0; font-weight: 760; letter-spacing: 0; line-height: 1.25; }
h1 { font-size: 32px; margin-bottom: 10px; }
h2 { font-size: 21px; margin-top: 32px; margin-bottom: 12px; padding-bottom: 7px; border-bottom: 1px solid var(--line); scroll-margin-top: 78px; }
p { margin: 0 0 12px; }
.subtitle { color: var(--muted); font-size: 16px; }
.notice {
  border: 1px solid var(--line);
  background: #ffffff;
  padding: 12px 14px;
  margin: 18px 0;
  line-height: 1.58;
  overflow-wrap: anywhere;
}
.report-nav {
  position: sticky;
  top: 0;
  z-index: 30;
  display: flex;
  align-items: center;
  gap: 8px;
  overflow-x: auto;
  white-space: nowrap;
  padding: 9px 10px;
  margin: 0 0 18px;
  border: 1px solid var(--line);
  background: #ffffff;
  box-shadow: 0 1px 4px rgb(24 33 47 / 0.08);
  scrollbar-width: thin;
}
.report-nav span {
  flex: 0 0 auto;
  color: var(--muted);
  font-size: 13px;
  font-weight: 700;
}
.report-nav a {
  flex: 0 0 auto;
  color: var(--ink);
  text-decoration: none;
  font-size: 13px;
  line-height: 1.2;
  padding: 5px 8px;
  border: 1px solid transparent;
}
.report-nav a:hover,
.report-nav a:focus {
  border-color: var(--line);
  outline: none;
}
.meta-grid {
  display: grid;
  grid-template-columns: 150px minmax(0, 1fr);
  gap: 0;
  border: 1px solid var(--line);
  margin: 14px 0 18px;
}
.meta-grid div { padding: 9px 11px; border-bottom: 1px solid var(--line); overflow-wrap: anywhere; }
.meta-grid div:nth-child(odd) { background: var(--soft); font-weight: 700; color: #243346; }
.meta-grid div:nth-last-child(-n+2) { border-bottom: 0; }
ul, ol { margin-top: 8px; padding-left: 22px; }
li { margin: 4px 0; }
.table-wrap { width: 100%; overflow-x: auto; margin: 12px 0 10px; }
table { width: 100%; border-collapse: collapse; table-layout: fixed; background: #ffffff; border: 1px solid var(--line); }
th, td {
  border: 1px solid var(--line);
  padding: 8px 9px;
  vertical-align: top;
  overflow-wrap: anywhere;
  word-break: break-word;
}
th { background: var(--soft); color: #1e2b39; font-weight: 730; text-align: left; }
tbody tr:nth-child(even) { background: #fbfcfd; }
td.center, th.center { text-align: center; }
.meaning {
  color: var(--muted);
  border-left: 3px solid var(--accent);
  padding-left: 10px;
  margin: 8px 0 18px;
}
.avoid-break { break-inside: avoid; page-break-inside: avoid; }
@page { size: A4; margin: 16mm 14mm; }
@media (max-width: 720px) {
  body { font-size: 14px; }
  .page { padding: 28px 18px 42px; }
  h1 { font-size: 27px; }
  h2 { font-size: 19px; scroll-margin-top: 86px; }
  .report-nav { margin-left: -18px; margin-right: -18px; border-left: 0; border-right: 0; }
  .meta-grid { grid-template-columns: 1fr; }
  .meta-grid div:nth-child(odd) { border-bottom: 0; }
  .meta-grid div:nth-last-child(-n+2) { border-bottom: 1px solid var(--line); }
  .meta-grid div:last-child { border-bottom: 0; }
}
@media print {
  html, body, .page { background: #ffffff !important; }
  body { font-size: 11px; line-height: 1.45; }
  .page { max-width: none; padding: 0; }
  .report-nav { display: none; }
  h1 { font-size: 24px; }
  h2 { font-size: 16px; break-after: avoid; page-break-after: avoid; }
  .table-wrap { overflow: visible; break-inside: auto; page-break-inside: auto; }
  table { table-layout: fixed; page-break-inside: auto; }
  tr { break-inside: avoid; page-break-inside: avoid; }
  th, td { padding: 5px 6px; }
}
"""
    parts = [
        "<!doctype html>",
        '<html lang="zh-Hans">',
        "<head>",
        '<meta charset="utf-8">',
        '<meta name="viewport" content="width=device-width, initial-scale=1">',
        f"<title>{html.escape(REPORT['title'])}</title>",
        f"<style>{css}</style>",
        "</head>",
        "<body><main class=\"page\">",
        "<section class=\"cover\">",
        f"<h1>{html.escape(REPORT['title'])}</h1>",
        f"<p class=\"subtitle\">{html.escape(REPORT['brand'])}｜{html.escape(REPORT['subtitle'])}</p>",
        "</section>",
        f"<p class=\"notice\">{html.escape(REPORT['notice'])}</p>",
    ]
    nav_items = [
        ("section-summary", "摘要"),
        ("section-method", "方法"),
        ("section-authority", "参考"),
        ("section-crawl", "抓取"),
        ("section-entity", "实体"),
        ("section-product", "产品"),
        ("section-questions", "问题"),
        ("section-technical", "技术"),
        ("section-scores", "评分"),
        ("section-onsite", "站内"),
        ("section-offsite", "站外"),
        ("section-risk", "风险"),
    ]
    parts.append('<nav class="report-nav" aria-label="报告章节菜单"><span>章节</span>')
    for anchor, label in nav_items:
        parts.append(f'<a href="#{anchor}">{html.escape(label)}</a>')
    parts.append("</nav>")
    parts.extend([
        '<h2 id="section-meta">诊断元信息</h2>',
        "<div class=\"meta-grid\">",
    ])
    for key, value in REPORT["meta"]:
        parts.append(f"<div>{html.escape(key)}</div><div>{html.escape(value)}</div>")
    parts.append("</div>")
    parts.append('<h2 id="section-summary">执行摘要</h2><ul>')
    for title, detail in REPORT["summary"]:
        parts.append(f"<li><strong>{html.escape(title)}</strong>：{html.escape(detail)}</li>")
    parts.append("</ul>")

    section_specs = [
        ("section-method", "方法依据与评分口径", "method_basis", REPORT["section_meanings"]["method_basis"]),
        ("section-authority", "权威参考映射", "authority_map", REPORT["section_meanings"]["authority_map"]),
        ("section-questions", "用户问题覆盖矩阵", "question_coverage", REPORT["section_meanings"]["question_coverage"]),
        ("section-baseline", "官网抓取与事实交叉验证", "baseline", REPORT["section_meanings"]["baseline"]),
        ("section-crawl", "官网抓取覆盖清单", "crawl_inventory", REPORT["section_meanings"]["crawl_inventory"]),
        ("section-entity", "品牌实体与事实档案", "entity_profile", REPORT["section_meanings"]["entity_profile"]),
        ("section-product", "产品服务与商业信息地图", "product_map", REPORT["section_meanings"]["product_map"]),
        ("section-technical", "技术可抓取与结构化数据诊断", "technical_audit", REPORT["section_meanings"]["technical_audit"]),
        ("section-competitors", "竞品与外部信源对标", "competitors", REPORT["section_meanings"]["competitors"]),
        ("section-scores", "GEO 特征评分", "scores", REPORT["section_meanings"]["scores"]),
        ("section-opportunities", "机会地图与优先级矩阵", "opportunities", REPORT["section_meanings"]["opportunities"]),
        ("section-onsite", "站内系统方案", "repairs", ""),
        ("section-offsite", "站外证据建设方案", "offsite_plan", ""),
        ("section-sources", "来源台账与待确认项", "sources", ""),
        ("section-risk", "风险、假设与待确认项", "risk_register", REPORT["section_meanings"]["risk_register"]),
    ]
    for anchor, title, key, meaning in section_specs:
        if key in {"sources", "risk_register"}:
            parts.append("<section class=\"avoid-break\">")
        parts.append(f'<h2 id="{anchor}">{html.escape(title)}</h2>')
        parts.append(render_html_table(HEADERS[key], REPORT[key]))
        if meaning:
            parts.append(f"<p class=\"meaning\">判断含义：{html.escape(meaning)}</p>")
        if key in {"sources", "risk_register"}:
            parts.append("</section>")
    parts.append('<h2 id="section-diagnostic">诊断问题池</h2><ol>')
    for question in REPORT["diagnostic_questions"]:
        parts.append(f"<li>{html.escape(question)}</li>")
    parts.append("</ol>")
    parts.append('<h2 id="section-conclusion">结论</h2>')
    parts.append(f"<p>{html.escape(REPORT['conclusion'])}</p>")
    parts.append("</main></body></html>")
    return "\n".join(parts)


def render_html_table(headers: list[str], rows: list[list[str]]) -> str:
    header_html = "".join(f"<th>{html.escape(item)}</th>" for item in headers)
    row_html = []
    for row in rows:
        row_html.append("<tr>" + "".join(f"<td>{html.escape(str(cell))}</td>" for cell in row) + "</tr>")
    return f"<div class=\"table-wrap\"><table><thead><tr>{header_html}</tr></thead><tbody>{''.join(row_html)}</tbody></table></div>"


def set_cell_text(cell, text: str, bold: bool = False, shade: str | None = None, font_size: float = 8.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_cell_margins(cell, top=90, start=90, bottom=90, end=90)
    if shade:
        set_cell_shading(cell, shade)


def cm_to_twips(width_cm: float) -> int:
    return int(width_cm / 2.54 * 1440)


def set_cell_width(cell, width_cm: float) -> None:
    cell.width = Cm(width_cm)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(cm_to_twips(width_cm)))


def set_fixed_table_layout(table, widths_cm: list[float]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    table_width = min(sum(widths_cm), DOCX_SAFE_TABLE_WIDTH_CM)
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(cm_to_twips(table_width)))
    tbl_indent = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_indent is None:
        tbl_indent = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_indent)
    tbl_indent.set(qn("w:type"), "dxa")
    tbl_indent.set(qn("w:w"), "0")
    tbl = table._tbl
    for existing in list(tbl.findall(qn("w:tblGrid"))):
        tbl.remove(existing)
    grid = OxmlElement("w:tblGrid")
    for width in widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(cm_to_twips(width)))
        grid.append(grid_col)
    tbl.insert(1, grid)
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def set_cell_margins(cell, top: int, start: int, bottom: int, end: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "D8E0E8")


def set_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "PingFang SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.22
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in [
        ("Title", 22, "18212F"),
        ("Heading 1", 15, "1D5D8C"),
        ("Heading 2", 12, "18212F"),
    ]:
        style = styles[name]
        style.font.name = "PingFang SC"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(7)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(table)
    font_size = 7.6 if len(headers) >= 8 else 8.0 if len(headers) >= 6 else 8.5
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, shade="F5F7FA", font_size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for idx, cell_text in enumerate(row):
            set_cell_text(cells[idx], cell_text, font_size=font_size)
    set_fixed_table_layout(table, column_widths(headers))
    doc.add_paragraph()


def column_widths(headers: list[str]) -> list[float]:
    presets = {
        tuple(HEADERS["method_basis"]): [2.2, 3.6, 5.0, 4.6],
        tuple(HEADERS["authority_map"]): [2.6, 4.2, 4.5, 4.5],
        tuple(HEADERS["question_coverage"]): [1.6, 5.6, 1.8, 3.6, 2.8],
        tuple(HEADERS["baseline"]): [2.0, 2.7, 3.6, 2.8, 1.7, 2.6],
        tuple(HEADERS["crawl_inventory"]): [1.8, 3.1, 1.6, 3.1, 3.0, 3.2],
        tuple(HEADERS["entity_profile"]): [2.0, 3.5, 2.8, 1.8, 5.1],
        tuple(HEADERS["product_map"]): [2.0, 3.2, 2.5, 1.8, 2.8, 3.5],
        tuple(HEADERS["technical_audit"]): [2.0, 1.6, 4.0, 3.5, 4.7],
        tuple(HEADERS["competitors"]): [2.1, 3.4, 3.2, 3.2, 3.3],
        tuple(HEADERS["scores"]): [2.0, 1.4, 7.3, 4.8],
        tuple(HEADERS["opportunities"]): [1.5, 1.8, 2.4, 1.1, 1.1, 1.1, 1.6, 3.6],
        tuple(HEADERS["repairs"]): [2.0, 2.7, 3.3, 2.1, 2.5, 2.9],
        tuple(HEADERS["offsite_plan"]): [2.0, 3.3, 3.2, 1.4, 2.2, 3.4],
        tuple(HEADERS["sources"]): [2.1, 2.8, 1.8, 2.9, 4.8],
        tuple(HEADERS["risk_register"]): [3.4, 1.5, 1.5, 2.5, 6.8],
        ("方案类型", "机会", "对应问题", "价值", "工作量", "优先级"): [2.0, 3.0, 4.0, 1.4, 1.4, 1.4],
        ("机会", "负责人建议", "验收标准"): [3.5, 3.0, 8.7],
    }
    return presets.get(tuple(headers), [DOCX_SAFE_TABLE_WIDTH_CM / len(headers)] * len(headers))


def add_docx_named_table(doc: Document, key: str) -> None:
    headers = HEADERS[key]
    rows = REPORT[key]
    if key == "opportunities":
        doc.add_paragraph("机会优先级", style="Heading 2")
        add_table(
            doc,
            ["方案类型", "机会", "对应问题", "价值", "工作量", "优先级"],
            [[row[0], row[1], row[2], row[3], row[4], row[5]] for row in rows],
        )
        doc.add_paragraph("责任人与验收", style="Heading 2")
        add_table(
            doc,
            ["机会", "负责人建议", "验收标准"],
            [[row[1], row[6], row[7]] for row in rows],
        )
        return
    add_table(doc, headers, rows)


def add_bullets(doc: Document, items: list[tuple[str, str]]) -> None:
    for title, detail in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(f"{title}：")
        run.bold = True
        run.font.name = "PingFang SC"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        paragraph.add_run(detail)


def render_docx() -> None:
    doc = Document()
    set_document_styles(doc)
    title = doc.add_paragraph(style="Title")
    title.add_run(REPORT["title"])
    subtitle = doc.add_paragraph()
    subtitle.add_run(f"{REPORT['brand']}｜{REPORT['subtitle']}").bold = True
    notice = doc.add_paragraph()
    notice.add_run("说明：").bold = True
    notice.add_run(REPORT["notice"])

    doc.add_heading("诊断元信息", level=1)
    meta_table = doc.add_table(rows=0, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_borders(meta_table)
    for key, value in REPORT["meta"]:
        cells = meta_table.add_row().cells
        set_cell_text(cells[0], key, bold=True, shade="F5F7FA")
        set_cell_text(cells[1], value)
    set_fixed_table_layout(meta_table, [3.0, 12.5])

    doc.add_heading("执行摘要", level=1)
    add_bullets(doc, REPORT["summary"])

    doc.add_heading("方法依据与评分口径", level=1)
    add_table(doc, HEADERS["method_basis"], REPORT["method_basis"])
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['method_basis']}")

    doc.add_heading("权威参考映射", level=1)
    add_docx_named_table(doc, "authority_map")
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['authority_map']}")

    doc.add_heading("用户问题覆盖矩阵", level=1)
    add_table(doc, HEADERS["question_coverage"], REPORT["question_coverage"])
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['question_coverage']}")

    doc.add_heading("官网抓取与事实交叉验证", level=1)
    add_docx_named_table(doc, "baseline")
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['baseline']}")

    doc.add_heading("官网抓取覆盖清单", level=1)
    add_docx_named_table(doc, "crawl_inventory")
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['crawl_inventory']}")

    doc.add_heading("品牌实体与事实档案", level=1)
    add_docx_named_table(doc, "entity_profile")
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['entity_profile']}")

    doc.add_heading("产品服务与商业信息地图", level=1)
    add_docx_named_table(doc, "product_map")
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['product_map']}")

    doc.add_heading("技术可抓取与结构化数据诊断", level=1)
    add_docx_named_table(doc, "technical_audit")
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['technical_audit']}")

    doc.add_heading("竞品与外部信源对标", level=1)
    add_docx_named_table(doc, "competitors")
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['competitors']}")

    doc.add_heading("GEO 特征评分", level=1)
    add_docx_named_table(doc, "scores")
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['scores']}")

    doc.add_heading("机会地图与优先级矩阵", level=1)
    add_docx_named_table(doc, "opportunities")
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['opportunities']}")

    doc.add_heading("站内系统方案", level=1)
    add_docx_named_table(doc, "repairs")

    doc.add_heading("站外证据建设方案", level=1)
    add_docx_named_table(doc, "offsite_plan")

    doc.add_heading("诊断问题池", level=1)
    for question in REPORT["diagnostic_questions"]:
        doc.add_paragraph(question, style="List Number")

    doc.add_heading("来源台账与待确认项", level=1)
    add_docx_named_table(doc, "sources")

    doc.add_heading("风险、假设与待确认项", level=1)
    add_docx_named_table(doc, "risk_register")
    doc.add_paragraph(f"判断含义：{REPORT['section_meanings']['risk_register']}")

    doc.add_heading("结论", level=1)
    doc.add_paragraph(REPORT["conclusion"])
    doc.save(DOCX_PATH)


def render_pdf() -> None:
    chrome = Path("/Applications/Google Chrome.app/Contents/MacOS/Google Chrome")
    if not chrome.exists():
        raise RuntimeError("未找到 Google Chrome，无法从 HTML 渲染 PDF。")
    with tempfile.TemporaryDirectory() as tmpdir:
        cmd = [
            str(chrome),
            "--headless=new",
            "--disable-gpu",
            "--no-sandbox",
            "--disable-background-networking",
            "--disable-sync",
            "--disable-extensions",
            "--disable-component-update",
            "--disable-default-apps",
            f"--user-data-dir={tmpdir}",
            f"--print-to-pdf={PDF_PATH}",
            "--no-pdf-header-footer",
            "--print-to-pdf-no-header",
            HTML_PATH.as_uri(),
        ]
        process = subprocess.Popen(cmd, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        try:
            process.wait(timeout=30)
        except subprocess.TimeoutExpired:
            process.terminate()
            try:
                process.wait(timeout=5)
            except subprocess.TimeoutExpired:
                process.kill()
        if not PDF_PATH.exists() or PDF_PATH.stat().st_size <= 1024:
            raise RuntimeError("PDF 未成功生成或文件过小。")


def review() -> None:
    checks: list[tuple[str, bool, str]] = []
    for path in [MD_PATH, HTML_PATH, DOCX_PATH, PDF_PATH]:
        checks.append((path.name, path.exists() and path.stat().st_size > 1024, "文件存在且非空"))
    html_text = HTML_PATH.read_text(encoding="utf-8")
    md_text = MD_PATH.read_text(encoding="utf-8")
    docx_doc = Document(DOCX_PATH) if DOCX_PATH.exists() else None
    docx_tables = len(docx_doc.tables) if docx_doc else 0
    max_docx_table_width = 0.0
    if docx_doc:
        for table in docx_doc.tables:
            if not table.rows:
                continue
            width = sum((cell.width.cm if cell.width else 0.0) for cell in table.rows[0].cells)
            max_docx_table_width = max(max_docx_table_width, width)
    pdf_text = ""
    if shutil.which("pdftotext") and PDF_PATH.exists():
        pdf_text = subprocess.run(
            ["pdftotext", str(PDF_PATH), "-"],
            check=False,
            text=True,
            stdout=subprocess.PIPE,
            stderr=subprocess.DEVNULL,
            timeout=15,
        ).stdout
    checks.extend(
        [
            ("HTML 白底", "background: #ffffff" in html_text and "color-scheme: light" in html_text, "白底与浅色模式存在"),
            ("HTML 防溢出", "overflow-wrap: anywhere" in html_text and "table-layout: fixed" in html_text, "表格与长文本防溢出规则存在"),
            ("HTML 固定菜单", "report-nav" in html_text and "position: sticky" in html_text and "section-risk" in html_text, "章节菜单固定跟随并包含核心锚点"),
            ("HTML 打印降级", ".report-nav { display: none; }" in html_text and "@media print" in html_text, "打印/PDF 时隐藏固定菜单"),
            ("Kami 版式细节", all(token in html_text for token in ["tbody tr:nth-child(even)", "@media (max-width: 720px)", "scrollbar-width: thin", "border: 1px solid var(--line);"]), "表格节奏、移动端菜单和边框细节存在"),
            ("中文字段", all(token in md_text for token in ["诊断元信息", "执行摘要", "方法依据与评分口径", "机会地图", "验收标准"]), "核心字段为中文简体"),
            ("问题覆盖", all(token in md_text for token in ["推荐", "购买决策", "场景解决"]), "九类意图覆盖矩阵存在"),
            ("完整性模块", all(token in md_text for token in ["权威参考映射", "官网抓取覆盖清单", "品牌实体与事实档案", "产品服务与商业信息地图", "技术可抓取与结构化数据诊断", "风险、假设与待确认项"]), "系统、详细、完整模块存在"),
            ("站内站外拆分", all(token in md_text for token in ["站内系统方案", "站外证据建设方案"]), "站内与站外方案均存在"),
            ("官网主线", all(token in md_text for token in ["官网抓取与事实交叉验证", "站内系统方案", "站外证据建设方案"]), "已切换为官网抓取与方案拆分章节"),
            ("格式一致", all(token in html_text for token in ["权威参考映射", "用户问题覆盖矩阵", "官网抓取覆盖清单", "站内系统方案", "站外证据建设方案", "诊断问题池"]), "HTML 与 Markdown 章节一致"),
            ("Word 表格结构", docx_tables >= 16, f"检测到 {docx_tables} 张表格"),
            ("Word 表格宽度", max_docx_table_width <= DOCX_SAFE_TABLE_WIDTH_CM + 0.1, f"最大表宽约 {max_docx_table_width:.1f}cm"),
            ("PDF 无本地路径页脚", "file:///" not in pdf_text and "Users/laoyao" not in pdf_text, "未检测到浏览器默认路径页脚"),
        ]
    )
    failed = [name for name, ok, _ in checks if not ok]
    lines = COPYRIGHT_COMMENT.rstrip().splitlines() + [
        "",
        "# 示例报告自检记录",
        "",
        "- 自检日期：2026-05-21",
        f"- 示例性质：{REPORT.get('example_nature', '官网抓取与公开事实交叉验证样例')}",
        "- 版式基准：kami 专业文档纪律 + 白底报告约束",
        "",
        "| 检查项 | 结果 | 说明 |",
        "| --- | --- | --- |",
    ]
    for name, ok, detail in checks:
        lines.append(f"| {name} | {'通过' if ok else '需修复'} | {detail} |")
    lines.append("")
    if failed:
        lines.append("## 待修复")
        for item in failed:
            lines.append(f"- {item}")
        REVIEW_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")
        raise RuntimeError("示例报告自检未通过：" + "、".join(failed))
    lines.append("## 自检结论")
    lines.append("四种格式均已生成；HTML/PDF 采用白底报告版式，HTML 增加固定跟随章节菜单、移动端单列元信息、打印菜单降级、斑马纹表格、固定布局与长文本换行；Word 使用标题样式、显式列宽、固定表格网格、表格边框和中文字段，密集 8 列机会表已按 Word 版拆成窄表。报告聚焦官网抓取、权威参考映射、事实交叉验证、完整性诊断、站内方案和站外方案。")
    lines.append("受本机缺少 LibreOffice 影响，未执行 DOCX 页面图片渲染检查；已完成结构与文件完整性检查。")
    REVIEW_PATH.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> None:
    MD_PATH.write_text(render_markdown(), encoding="utf-8")
    HTML_PATH.write_text(render_html(), encoding="utf-8")
    render_docx()
    render_pdf()
    review()


if __name__ == "__main__":
    main()
